from __future__ import annotations

import argparse
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SKILL_DIR = Path(os.environ["REMIN_DOCUMENT_SKILL_DIR"])
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


INK = RGBColor(29, 45, 68)
BLUE = RGBColor(74, 113, 177)
MUTED = RGBColor(92, 101, 116)
HEADER_FILL = "EEF2F8"


def set_font(run, *, size: float, bold: bool = False, color: RGBColor = INK) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    properties.append(shade)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, INK),
    ):
        style = document.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build(output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("拾忆阶段0测试资料 · DOCX"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("完全离线 · 源文件只读"), size=9, color=MUTED)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("归航计划项目总结"), size=23, bold=True, color=INK)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("用于中文检索、问答、抽取和原文定位的确定性测试资料"), size=12, color=MUTED)

    for label, value in (
        ("项目编号", "GH-2025-017"),
        ("项目负责人", "林晓岚"),
        ("评审日期", "2025-11-18"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(f"{label}："), size=11, bold=True)
        set_font(paragraph.add_run(value), size=11)

    document.add_heading("1. 项目概览", level=1)
    document.add_paragraph(
        "归航计划于2025年3月启动，目标是在完全离线的Windows电脑上帮助用户找回散落资料。项目不连接Ollama，源文件始终只读，索引和日志仅保存在拾忆管理的本地目录。"
    )
    document.add_paragraph(
        "本轮预算为286,500元，其中硬件适配与性能测试占86,000元。负责人林晓岚要求所有事实性回答都能回到原文件、段落、页码或单元格。"
    )

    document.add_heading("2. 检索方案", level=1)
    document.add_paragraph(
        "文件名、SQLite FTS5全文和向量语义结果分别召回，再使用倒数排名融合合并候选。评测基线固定RRF k=60；语义通道不可用时，系统保留文件名和全文结果，不阻塞基础模式。"
    )

    table = document.add_table(rows=1, cols=4)
    headers = ["检查项", "目标", "负责人", "状态"]
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, HEADER_FILL)
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10, bold=True)
    rows = [
        ("中文已知目标搜索", "Recall@10不低于90%", "林晓岚", "进行中"),
        ("事实性问答", "引用覆盖率100%", "周予安", "待验证"),
        ("源文件保护", "E2E前后SHA-256不变", "陈默", "已通过"),
        ("无证据问题", "拒答率不低于95%", "周予安", "待验证"),
    ]
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cells[index].paragraphs[0].runs:
                set_font(run, size=10)
    apply_table_geometry(table, [1800, 3300, 1500, 2760], table_width_dxa=9360, indent_dxa=120)

    document.add_heading("3. 明确结论", level=1)
    document.add_paragraph(
        "首次验收会议定于2025年11月18日14:30。若可用内存低于安全阈值，系统从full降为balanced，再降为core；任何降级都不能取消目录授权、只读、Schema和引用检查。"
    )
    document.core_properties.title = "归航计划项目总结"
    document.core_properties.subject = "拾忆阶段0确定性测试资料"
    document.core_properties.author = "Remin Test Corpus"
    document.core_properties.keywords = "拾忆,归航计划,GH-2025-017,RRF"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.output)


if __name__ == "__main__":
    main()
